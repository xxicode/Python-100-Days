"""
读取Word文件

Version: 0.1
Author: 骆昊
Date: 2018-03-26
"""


from docx import Document

doc = Document('./res/用函数还是用复杂的表达式.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
content = [para.text for para in doc.paragraphs]
print(''.join(content))
